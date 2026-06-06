"""Generate methodology.docx for management and internal review teams.

Run: uv run python scripts/build_methodology_doc.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_PATH = Path(__file__).resolve().parent.parent / "methodology.docx"

PRIMARY = RGBColor(0x14, 0x3D, 0x6B)
ACCENT = RGBColor(0x4A, 0x6F, 0xA5)
MUTED = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x21, 0x6E, 0x39)
AMBER = RGBColor(0x9C, 0x6B, 0x00)


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY


def _add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = ""
        para = hdr_cells[idx].paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[idx], "143D6B")
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def _add_callout(doc: Document, label: str, body: str, *, color: RGBColor = PRIMARY, fill: str = "EAF1F8") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill)
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = color
    para.add_run(body)
    doc.add_paragraph()


def _add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def _add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    for line in text.splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)


def build_document() -> Document:  # noqa: PLR0915 - long but linear document builder
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ====================================================================
    # COVER
    # ====================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Credit Risk Scoring & Cutoff Optimization")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Methodology Document")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = ACCENT

    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aud_run = audience.add_run("Audience: Senior Management and Internal Review Teams")
    aud_run.italic = True
    aud_run.font.size = Pt(12)
    aud_run.font.color.rgb = MUTED

    doc.add_paragraph()
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Light List Accent 1"
    meta_rows = [
        (
            "Document scope",
            "End-to-end methodology for credit risk scoring, cutoff optimization, out-of-time validation, and portfolio allocation",
        ),
        ("Owner", "Risk Analytics"),
        ("Status", "For management and internal review"),
        ("Review cadence", "On material methodology change; annual otherwise"),
        ("Primary risk metric", "b2_ever_h6 — annualised, exposure-weighted 6-month vintage delinquency rate"),
        ("Decision artefact", "Pareto-efficient menu of cutoff policies with bootstrap CIs and out-of-time validation"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ====================================================================
    # 1. EXECUTIVE SUMMARY
    # ====================================================================
    _add_heading(doc, "1. Executive Summary", 1)
    _add_paragraph(
        doc,
        "This document explains the methodology behind the credit-scoring and cutoff-optimization "
        "platform: what it does, the analytical choices behind it, and why each choice was made over "
        "the alternatives. It is written for senior management and internal review audiences and "
        "favours rationale, formal definitions, and control evidence over implementation detail.",
    )
    _add_paragraph(
        doc,
        "The platform answers a single business question: given the credit scores of incoming "
        "applications, what acceptance policy maximises booked production while keeping the "
        "annualised 6-month vintage delinquency rate (b2_ever_h6) within an agreed risk budget? "
        "The five-step answer is:",
    )
    _add_numbered(
        doc,
        [
            "Discretise scores into an N-dimensional grid of bins (typically internal score x external bureau score, optionally x income).",
            "Estimate risk per cell for both booked applications (observed) and score-rejected applications (inferred via a trained model and reject-inference correction).",
            "Solve a Mixed-Integer Linear Program (MILP) that selects which grid cells to accept, subject to a risk budget and monotonicity constraints.",
            "Build a Pareto frontier of risk vs. production at ~50 risk targets, then highlight three operating points — pessimistic, base, optimistic — derived from the configured optimum_risk and risk_step.",
            "Validate the chosen cutoffs on a holdout out-of-time period (Recent Monitoring or MR) and report stability (PSI/CSI), uncertainty (1,000-replicate bootstrap CIs), and sensitivity to risk perturbations of plus/minus 5%, 10%, 20%.",
        ],
    )
    _add_callout(
        doc,
        "Three guiding principles",
        "Every methodological choice in this document is driven by: (1) regulatory defensibility — "
        "decisions must be auditable and reproducible, with seeded randomness and immutable "
        "configuration; (2) statistical robustness — uncertainty is quantified rather than hidden, "
        "and selection bias is corrected explicitly; (3) operational realism — the policy must be "
        "implementable as bin-level cutoffs in front-line systems, not as opaque ML scores.",
    )
    _add_heading(doc, "1.1 What management gets at the end of a run", 2)
    _add_bullets(
        doc,
        [
            "A self-contained HTML report per segment with the proposed cutoff grid, risk and production for actual / swap-in / swap-out / optimum, MR validation, PSI/CSI, and sensitivity panels.",
            "An Excel workbook (consolidated_risk_production.xlsx) whose first three sheets form a trust layer — an executive summary (KPI cards with bootstrap CIs and a plain-language recommendation), a Validation & Governance sheet (data-snapshot provenance, assumption tiers, per-segment reproducibility and stability), and an Out-of-time Validation sheet — followed by per-segment risk-production sheets; designed to be circulated in committee.",
            "A reproducible TOML configuration that fully describes the run; reruns with the same TOML and the same data produce byte-identical outputs.",
            "An efficient frontier per segment (efficient_frontier_{scenario}.csv) and an optional global allocation (allocation_results.csv plus a Markdown narrative) that distributes a portfolio-wide risk budget across segments.",
        ],
    )

    # ====================================================================
    # 2. BUSINESS CONTEXT
    # ====================================================================
    _add_heading(doc, "2. Business Context and Objectives", 1)
    _add_paragraph(
        doc,
        "Credit scoring policies are the primary risk-control lever in unsecured consumer lending. "
        "A cutoff that is too tight forfeits profitable volume; one that is too loose erodes margin "
        "through delinquency. The objective of this platform is to make that trade-off explicit, "
        "evidence-based, and reproducible across portfolio segments.",
    )
    _add_paragraph(
        doc,
        "Historically, cutoffs were set by expert judgment with periodic vintage reviews. That "
        "approach has three structural weaknesses: (i) it produces a single point estimate, hiding "
        "the full risk-production trade-off; (ii) it is hard to validate out-of-time because the "
        "decision logic is implicit; (iii) it scales poorly across many segments. The platform "
        "addresses each in turn.",
    )
    _add_heading(doc, "2.1 Specific business objectives", 2)
    _add_table(
        doc,
        ["Objective", "How the platform delivers it"],
        [
            [
                "Maximise booked production for a given risk appetite",
                "MILP objective max sum oa_amt_h0[i] * x[i] subject to a linearised risk-budget constraint.",
            ],
            [
                "Provide a Pareto-efficient menu of policies, not a single answer",
                "Sweep ~50 risk targets, deduplicate by acceptance mask, filter to non-dominated points (lower-or-equal risk and higher-or-equal production with strict inequality).",
            ],
            [
                "Honour business constraints",
                "Per-segment min_risk / max_risk / min_production, swap-in production-share cap, swap-in risk cap, and forced-accept / forced-reject cells, all enforced inside the MILP.",
            ],
            [
                "Produce operationally simple policies",
                "Marginal monotonicity constraints force a staircase pattern: 'accept everything to the left of and below this line per dimension'.",
            ],
            [
                "Deliver an auditable trail",
                "Versioned TOML config, seeded randomness (numpy.random.RandomState(42), Optuna seed=42), per-bin risk_source column logging which evidence drove each MR estimate, ri_optimizer_results.csv recording every parameter pair evaluated.",
            ],
            [
                "Detect drift before it becomes a production incident",
                "PSI/CSI between main and MR periods, monthly SPC anomaly detection on approval rate / production / risk, MR degradation ratio for RI parameters.",
            ],
        ],
    )
    _add_heading(doc, "2.2 Scope and out-of-scope", 2)
    _add_table(
        doc,
        ["In scope", "Out of scope"],
        [
            ["Acceptance cutoffs on score-driven decisions", "Pricing decisions (rate / term / fees)"],
            [
                "Risk inference for score-rejected (repesca) applications",
                "Underwriting policy rules unrelated to score (fraud, KYC, documentation)",
            ],
            [
                "Out-of-time validation on a recent monitoring (MR) period",
                "Through-the-cycle macroeconomic stress modelling",
            ],
            [
                "Multi-segment allocation against a global portfolio risk budget",
                "Collections strategy after origination",
            ],
            [
                "Bin-level monotonic acceptance grids",
                "Continuous score-cutoff curves (deliberately rejected — see Section 6)",
            ],
            [
                "Annualised risk metrics on b2_ever_h6 / b2_ever_h3",
                "Lifetime PD/LGD modelling for IFRS 9 / CECL provisioning",
            ],
        ],
    )

    # ====================================================================
    # 3. PIPELINE PHASES
    # ====================================================================
    _add_heading(doc, "3. Methodology at a Glance", 1)
    _add_paragraph(
        doc,
        "The pipeline runs in eight phases. Each phase produces auditable artefacts (CSVs, model "
        "files, or HTML reports) that feed the next phase, so any single decision can be traced "
        "backwards to the inputs that justified it.",
    )
    _add_table(
        doc,
        ["#", "Phase", "Purpose", "Why it matters"],
        [
            [
                "1",
                "Configuration",
                "Load and validate parameters via Pydantic schema (PreprocessingSettings)",
                "Fail fast on bad inputs; every run is fully described by a TOML file. Schema validators check >= 2 variables, monotonic bin edges, parseable dates, paired MR period, range constraints.",
            ],
            [
                "2",
                "Data loading",
                "Read SAS files (.sas7bdat), standardise columns to lowercase_underscore",
                "Single source of truth for downstream steps; deterministic naming. In batch mode data is loaded once and shared across segments.",
            ],
            [
                "3",
                "Preprocessing",
                "DQ checks, segment/date filtering, binning, stress factor, transformation rate (tasa_fin)",
                "Quality gates before any modelling; explicit handling of selection bias via stress and tasa_fin.",
            ],
            [
                "4",
                "Inference (modelling)",
                "Train risk model with 5-fold CV, 1-SE selection, monotonic constraints; SHAP and per-cell CIs as side outputs",
                "Disciplined model choice that resists overfitting and respects business priors.",
            ],
            [
                "5",
                "Optimization",
                "MILP over monotonic cutoff combinations at ~50 risk targets; Pareto filter; optional GA fallback for N>2 if infeasible",
                "Globally optimal policies under explicit constraints — not greedy heuristics.",
            ],
            [
                "6",
                "Scenario analysis",
                "Pessimistic / base / optimistic operating points with 1,000-replicate bootstrap CIs",
                "Decisions framed under uncertainty, not as point estimates.",
            ],
            [
                "7",
                "Sensitivity & marginal impact",
                "Cell-level flip thresholds at +/- 5%, 10%, 20% perturbations; analytical O(N) marginal impact per cell",
                "Quantifies how robust the chosen policy is to small risk shocks.",
            ],
            [
                "8",
                "MR validation, stability, trends",
                "Out-of-time validation, PSI/CSI vs main period, monthly SPC anomaly detection, audit table per record",
                "Detects drift before it becomes a problem in production; provides record-level audit trail.",
            ],
        ],
    )
    _add_callout(
        doc,
        "Idempotency",
        "Reruns with the same TOML and the same source data produce byte-identical outputs. All "
        "randomness is seeded (numpy seed 42, Optuna seed 42, scipy MILP is deterministic). "
        "Bootstrap, CV folds, RI optimisation, and trend SPC all derive from these seeds.",
    )

    # ====================================================================
    # 4. DATA POPULATIONS
    # ====================================================================
    _add_heading(doc, "4. Data Populations and Why They Are Treated Differently", 1)
    _add_paragraph(
        doc,
        "Three application populations co-exist in the data and must be handled separately because "
        "they carry different information about creditworthiness:",
    )
    _add_table(
        doc,
        ["Population", "Definition (reject_reason)", "Sample size signal", "Treatment in the pipeline"],
        [
            [
                "Booked",
                "Approved and disbursed (no rejection reason)",
                "Largest, fully observed",
                "Observed outcomes (todu_30ever_h6 and todu_amt_pile_h6) used directly; aggregated and annualised per cell.",
            ],
            [
                "Score-rejected (repesca)",
                "'09-score' — rejected because the score fell below the cutoff",
                "Material; the population whose acceptance is being optimised",
                "Outcomes unobserved; estimated by trained model -> stress factor -> reject-inference parceling -> tasa_fin.",
            ],
            [
                "Policy-rejected",
                "'08-other' — rejected for fraud, KYC, documentation, policy rules",
                "Variable",
                "Excluded from optimisation entirely. Their rejection has no causal link to the score cutoff and including them would attribute non-credit rejections to credit decisioning.",
            ],
        ],
    )
    _add_callout(
        doc,
        "Why excluding policy rejections is essential",
        "Mixing policy rejections into the optimisation pool would (i) bias the learned acceptance "
        "rate denominator used by reject inference, (ii) inflate the apparent rejection rate of "
        "low-score bins, and (iii) make the policy uninterpretable because two distinct decisions "
        "(score cutoff vs. policy rule) would be conflated. The reject_include_all_rejections flag "
        "is deprecated and ignored — acceptance rates are always computed score-only because the "
        "swap-in population is solely score-rejected, so setting it has no effect (it logs a "
        "one-time warning).",
    )
    _add_heading(doc, "4.1 The audit table", 2)
    _add_paragraph(
        doc,
        "For each scenario, the platform produces a record-level audit table classifying every "
        "application against the proposed cutoff:",
    )
    _add_table(
        doc,
        ["Classification", "Origin", "Outcome under proposed policy"],
        [
            ["keep", "Booked", "Still accepted -- core portfolio"],
            ["swap_out", "Booked", "Rejected -- the policy is tightening here"],
            ["swap_in", "Score-rejected (09-score)", "Now accepted -- policy is loosening (repesca)"],
            ["rejected", "Score-rejected (09-score)", "Still rejected"],
            ["rejected_other", "Policy-rejected (08-other)", "Always rejected -- not in scope"],
        ],
    )
    _add_paragraph(
        doc,
        "This table is the single source of truth for downstream impact analysis: it is what "
        "underwriting committees, audit, and regulators see when they ask 'who exactly does this "
        "policy change affect?'",
    )

    # ====================================================================
    # 5. RISK METRIC
    # ====================================================================
    _add_heading(doc, "5. The Risk Metric: b2_ever_h6", 1)
    _add_paragraph(
        doc,
        "The primary risk indicator is b2_ever_h6, the annualised 6-month vintage rate of accounts "
        "ever 30+ days past due, weighted by exposure:",
    )
    _add_formula(doc, "b2_ever_h6 = multiplier x todu_30ever_h6 / todu_amt_pile_h6")
    _add_paragraph(doc, "Where:")
    _add_bullets(
        doc,
        [
            "todu_30ever_h6 is the count of applications that ever reached 30+ days past due within a 6-month observation horizon.",
            "todu_amt_pile_h6 is the total cumulative exposure amount within the same horizon (the denominator that turns counts into an exposure-weighted rate).",
            "multiplier is an annualisation constant (default 7.0 for H6, 4.0 for H3). It compresses a 6-month observation into a 12-month equivalent so the metric reads like an annual rate.",
        ],
    )
    _add_paragraph(
        doc,
        "The metric is clipped at zero (it cannot be negative) and is NaN when exposure is zero, "
        "which signals an empty cell rather than zero risk.",
    )

    _add_heading(doc, "5.1 Why a 6-month horizon", 2)
    _add_bullets(
        doc,
        [
            "Six months is long enough for early-life delinquency to crystallise on consumer credit products: most first defaults occur within months 2-5 after origination.",
            "Short enough to allow timely reaction: a 12-month horizon would delay any policy correction by 6 months relative to H6.",
            "Aligns with internal vintage reporting and external benchmarks, so cutoff decisions can be defended in management committees and against regulatory peer benchmarks.",
            "A complementary 3-month metric (b2_ever_h3) is computed for early warning use when the MR window has not fully matured to 6 months — see Section 11.2 on H3 -> H6 extrapolation.",
        ],
    )
    _add_heading(doc, "5.2 Why exposure-weighted rather than account-count", 2)
    _add_paragraph(
        doc,
        "Counting accounts treats a 1,000 EUR overdraft and a 30,000 EUR personal loan as equally "
        "costly, which misrepresents economic loss. Weighting the rate by exposure (todu_amt_pile_h6) "
        "ensures the metric tracks the financial impact the business actually carries:",
    )
    _add_formula(
        doc,
        "exposure-weighted rate = sum(exposure_i * default_i) / sum(exposure_i)",
    )
    _add_paragraph(
        doc,
        "This is the only formulation that lines up with provisioning math (where loss is "
        "exposure-at-default times PD times LGD), which means the cutoff optimisation directly "
        "moves the variable that drives P&L.",
    )
    _add_heading(doc, "5.3 Why annualisation", 2)
    _add_paragraph(
        doc,
        "Observation periods can vary across runs (a quarterly refresh runs over a 3-month window; "
        "a half-year refresh over 6 months). A constant annual coefficient normalises every metric "
        "to a 12-month equivalent:",
    )
    _add_formula(doc, "annual_coef = 12 / n_months_in_observation_period")
    _add_paragraph(
        doc,
        "All flow indicators (todu_30ever_h6, todu_amt_pile_h6, oa_amt_h0) are scaled by annual_coef "
        "as part of grid aggregation — the per-bin sums are multiplied by annual_coef (identical to "
        "scaling each row first) — so risk and production are comparable across runs of different "
        "lengths and across segments observed over different windows.",
    )

    # ====================================================================
    # 6. BINNING
    # ====================================================================
    _add_heading(doc, "6. N-Dimensional Binning of Scores", 1)
    _add_paragraph(
        doc,
        "Scores are continuous, but underwriting policies must be expressed as a small set of "
        "human-readable cutoffs. The platform discretises each score into ordered bins and "
        "operates on the resulting N-dimensional grid (typically internal x bureau, optionally x "
        "income).",
    )
    _add_heading(doc, "6.1 Why discretise at all", 2)
    _add_bullets(
        doc,
        [
            "Operational simplicity: front-line systems and policy committees reason in terms of grid cells, not continuous thresholds. A '10x10 grid with 6 accepted bins on the diagonal' is reviewable; a continuous bivariate cutoff curve is not.",
            "Statistical stability: aggregating into bins reduces noise and makes per-cell estimates defensible. A bin with 500 booked records produces a stable per-cell rate; a 0.01-wide score interval almost never does.",
            "Auditability: a binned grid produces a small, inspectable acceptance mask (the cells_accepted CSV) rather than an opaque score-cutoff curve.",
            "Monotonicity is naturally expressible: the MILP needs a discrete adjacency structure to enforce 'safer cell accepted implies any safer-still cell accepted', which is straightforward on a grid and awkward in a continuous space.",
            "Interpretability under model changes: when a score model is recalibrated, bin boundaries need to be re-learned but the policy structure (a staircase over the grid) survives. A continuous cutoff would silently shift.",
        ],
    )
    _add_heading(doc, "6.2 Two binning methods, and when each is used", 2)
    _add_table(
        doc,
        ["Method", "How it works", "When and why"],
        [
            [
                "Quantile (default)",
                "Equal-count bins learned on the demand population from the empirical distribution of the source column. Bin edges are quantile cut points so each bin holds approximately the same number of observations.",
                "All models — the only supported learned-edge method. Leakage-free (the boundaries ignore the target) and statistically stable in every cell.",
            ],
            [
                "Optimization (deprecated)",
                "Formerly fit a sklearn DecisionTreeRegressor weighted by production on the risk target and used its split points. Now deprecated and silently converted to quantile at runtime (logs a warning).",
                "Do not use. It learned bin boundaries from the same risk target the optimizer later maximises (target leakage). For established legacy tiers, set explicit bin_edges instead.",
            ],
        ],
    )
    _add_callout(
        doc,
        "Why quantile is the only learned-edge method",
        "Quantile bins are leakage-free by construction — they ignore the target, so they cannot "
        "overfit it. The earlier 'optimization' method learned bin edges from the very risk target "
        "the optimizer then maximises, leaking it into the grid; it is deprecated and now falls back "
        "to quantile with a warning. For established legacy tiers, set explicit bin_edges. If "
        "quantile bins make an extra dimension (e.g. income) look flat, that dimension genuinely "
        "lacks discriminating power rather than something to engineer around with supervised splits.",
    )
    _add_heading(doc, "6.3 Bin direction and risk ordering", 2)
    _add_paragraph(
        doc,
        "Each variable has an associated direction that encodes the risk ordering of its bins. This "
        "is what allows the MILP to enforce monotonicity correctly:",
    )
    _add_table(
        doc,
        ["Direction value", "Interpretation", "Typical example"],
        [
            [
                "-1 (inverted)",
                "Higher bin index = safer",
                "Internal score, bureau score — higher score is better creditworthiness",
            ],
            ["+1 (default)", "Higher bin index = riskier", "Loan amount, term — larger / longer is generally riskier"],
        ],
    )
    _add_paragraph(
        doc,
        "Directions can be set explicitly in the directions section of the config, or inferred from "
        "the data by computing the sign of the rank correlation between bin index and risk in the "
        "booked population. Manual override is preferred in committee-reviewed segments because it "
        "removes a hidden dependency on the training sample.",
    )
    _add_heading(doc, "6.4 Global bin learning across segments (batch mode)", 2)
    _add_paragraph(
        doc,
        "When several segments share an underlying score (e.g. all consumer-loan segments use the "
        "same internal score), bin edges are learned once on the combined population and reused. "
        "This means: (i) a cell labelled 'octroi=3, efx=5' has the same definition across all "
        "segments, so cross-segment comparisons are valid; (ii) the consolidated report can stack "
        "per-segment grids without grid mismatches; (iii) the global allocation MILP can compare "
        "frontiers without renormalising. Per-segment overrides remain available for cases where a "
        "segment's score distribution genuinely differs.",
    )

    # ====================================================================
    # 7. RISK MODEL
    # ====================================================================
    _add_heading(doc, "7. The Risk Model and How It Is Selected", 1)
    _add_paragraph(
        doc,
        "A regression model is trained on booked applications to predict b2_ever_h6 as a function "
        "of the score variables (specifically inference_variables, a subset of variables). Its "
        "predictions are used to (a) smooth the booked-only cells where direct estimates are noisy "
        "and (b) predict risk for score-rejected (repesca) cells where outcomes are unobserved.",
    )
    _add_heading(doc, "7.1 Candidate models and why this set", 2)
    _add_paragraph(
        doc,
        "The selection set deliberately spans a range of bias-variance trade-offs and modelling "
        "assumptions, so the chosen model reflects what the data actually supports rather than "
        "what was tried:",
    )
    _add_table(
        doc,
        ["Family", "Models", "Why included"],
        [
            [
                "Linear / regularised linear",
                "LinearRegression, Ridge, Lasso, ElasticNet",
                "Highly interpretable baselines; strong regularisation; robust to small samples; defensible when the risk surface is approximately additive.",
            ],
            [
                "GLM / specialised",
                "TweedieGLM, HurdleRegressor",
                "Tailored to zero-inflated, exposure-weighted insurance / credit-style targets where many cells have zero observed default.",
            ],
            [
                "Tree-based with monotonic constraints",
                "XGBoost, LightGBM (tuned via Optuna)",
                "Captures non-linear interactions while respecting business priors via per-feature monotonic constraints derived from the configured directions.",
            ],
        ],
    )
    _add_heading(doc, "7.2 Why monotonicity is enforced on tree models", 2)
    _add_paragraph(
        doc,
        "Without monotonic constraints, gradient-boosted trees can produce locally non-monotone "
        "risk surfaces — small score improvements can appear to raise predicted risk because of "
        "noise. That is operationally indefensible: any policy derived from such a surface would "
        "fail face validity in committee and create the appearance of arbitrary, even "
        "discriminatory, treatment of borderline applicants.",
    )
    _add_paragraph(
        doc,
        "The platform translates the configured directions dict into monotone_constraints for "
        "XGBoost / LightGBM (one entry per inference variable, +1 or -1 according to the variable's "
        "risk direction). These constraints are honoured by the solver during tree growth, so the "
        "resulting model is monotone by construction — not 'monotone in expectation' or 'monotone "
        "after smoothing'.",
    )
    _add_heading(doc, "7.3 Cross-validation discipline", 2)
    _add_paragraph(
        doc,
        "Model selection is a notorious source of optimistic bias. The platform applies four "
        "specific disciplines to keep selection honest:",
    )
    _add_numbered(
        doc,
        [
            "Five-fold CV with the split performed on the raw, unaggregated data; aggregation and outlier filtering are then applied independently to train and validation folds. Aggregating before splitting would leak target information into the validation set because the same per-cell mean would appear on both sides.",
            "Models and feature sets are ranked by weighted CV RMSE, with the one-standard-error rule used to pick the simplest candidate within one SE of the lowest mean RMSE. This is the standard defence against over-fitting to fold idiosyncrasies and biases the methodology toward parsimonious models.",
            "After Optuna selects hyperparameters, candidates are re-evaluated on fresh folds; with sufficient raw data, a nested holdout split provides a less biased generalisation estimate, separating tuning noise from selection noise.",
            "A single weight precedence is used for tuning, selection, cell-level CIs, and final fitting: todu_amt_pile_h6 (exposure) -> oa_amt_h0 (production) -> n_observations (count). This eliminates the common pitfall of selecting on one metric and reporting on another.",
        ],
    )
    _add_heading(doc, "7.4 Why the one-standard-error rule rather than 'best mean RMSE'", 2)
    _add_paragraph(
        doc,
        "Fold-level RMSE has its own variance. Selecting strictly on the minimum mean inflates "
        "complexity by chasing noise: across re-runs with shuffled folds, the 'winning' model "
        "frequently changes. The 1-SE rule trades a small amount of expected accuracy for a "
        "substantial gain in stability across re-runs and across periods — a property reviewers "
        "and management committees consistently prefer because it makes the methodology auditable "
        "rather than reactive.",
    )
    _add_heading(doc, "7.5 Continuous-impact sanity check", 2)
    _add_paragraph(
        doc,
        "After model selection, a secondary _estimate_continuous_impact analysis verifies the "
        "linear relationship between continuous production volume and risk. Because this runs on "
        "heavily compressed, post-aggregated groups (where N is tiny), it uses Leave-One-Out CV "
        "paired with cross_val_predict — the mathematically stable choice for a small-sample "
        "global R^2 trend estimate. This is a sanity check, not a selection criterion: a failure "
        "here flags that the model is fitting noise, not signal.",
    )
    _add_heading(doc, "7.6 Diagnostics shipped with every model", 2)
    _add_bullets(
        doc,
        [
            "SHAP feature importance saved as an image and as raw values, so the relative weight of each variable can be discussed in committee.",
            "Per-cell prediction confidence intervals (cell_level_ci.csv) computed by aggregating per-fold predictions and applying a Student-t critical value for small fold counts, normal otherwise. Cells observed in fewer than 2 folds receive NaN intervals — surfaced rather than hidden.",
            "Model metadata (CV scores, features, hyperparameters, training timestamp) saved alongside the joblib model file, so any deployed model can be traced back to its training run.",
        ],
    )

    # ====================================================================
    # 8. REJECT INFERENCE
    # ====================================================================
    _add_heading(doc, "8. Reject Inference: Correcting for Selection Bias", 1)
    _add_paragraph(
        doc,
        "Score-rejected applications were never disbursed, so their outcomes are not directly "
        "observable. Yet they are precisely the population whose acceptance is being evaluated. "
        "Two complementary mechanisms address this gap: a stress factor that scales repesca risk "
        "predictions toward observed tail risk, and a parceling correction that further uplifts "
        "risk based on per-bin acceptance rates.",
    )
    _add_heading(doc, "8.1 Stress factor", 2)
    _add_paragraph(
        doc,
        "The stress factor scales repesca risk predictions by the ratio of risk in the worst-"
        "performing 5% of booked applications to the overall booked rate:",
    )
    _add_formula(doc, "stress_factor = worst_5pct_bad_rate / overall_bad_rate")
    _add_formula(doc, "b2_ever_h6_repesca = stress_factor x model.predict(X)")
    _add_table(
        doc,
        ["Mode", "Behaviour", "When to use it"],
        [
            [
                "global (default)",
                "Single scalar across the grid",
                "Conservative and easy to explain in committee; the safe default.",
            ],
            [
                "per_bin",
                "Separate factor per cell; falls back to global if cell has fewer than 20 booked records",
                "Use when score variables span very different risk profiles, so tail-risk concentration is heterogeneous across the grid.",
            ],
            [
                "disabled",
                "Stress factor = 1.0",
                "Use when reject-inference parceling is active (avoids double-counting selection bias) or when the model is trained on through-the-door data.",
            ],
        ],
    )
    _add_heading(doc, "8.2 Reject-inference parceling: three functional forms", 2)
    _add_paragraph(
        doc,
        "After the stress adjustment, an optional parceling correction further uplifts repesca "
        "risk based on per-bin acceptance rates. The acceptance rate per cell is:",
    )
    _add_formula(doc, "acceptance_rate = n_booked / (n_booked + n_score_rejected)")
    _add_paragraph(
        doc,
        "Only score rejections (09-score) count in the denominator by default; policy rejections "
        "(08-other) are excluded because their rejection is unrelated to the credit-score cutoff. "
        "Three functional forms are offered because no single form fits every portfolio:",
    )
    _add_table(
        doc,
        ["Method", "Formula", "Risk shape", "When to choose it"],
        [
            [
                "Linear (default)",
                "1 + uplift * (1 - rate)",
                "Steady, interpretable penalty: 1.0x at full acceptance, 1 + uplift at zero acceptance",
                "General-purpose default; easiest to defend in committee because the relationship is obviously linear.",
            ],
            [
                "Power",
                "(1 / max(rate, 0.01)) ^ uplift",
                "Super-linear at low acceptance: rate=0.5 -> 2x at uplift=1; rate=0.1 -> 10x at uplift=1",
                "When risk is heavy-tailed and concentrates in highly-rejected bins. Power=1.0 reproduces the classic 1/a model.",
            ],
            [
                "Sigmoid",
                "1 + uplift / (1 + exp(10 * (rate - 0.5)))",
                "Smooth S-curve with steep transition near 50% acceptance",
                "When the risk-selectivity relationship saturates — very selective bins are not infinitely riskier, and fully accepting bins still have some baseline uncertainty.",
            ],
        ],
    )
    _add_paragraph(doc, "Numerical examples at uplift_factor = 1.5 (raw multipliers, before the floor of 1.0 and the default cap of 3.0):")
    _add_table(
        doc,
        ["Acceptance rate", "Linear multiplier", "Power multiplier", "Sigmoid multiplier"],
        [
            ["1.00", "1.00x", "1.00x", "approx 1.00x"],
            ["0.70", "1.45x", "approx 1.71x", "approx 1.18x"],
            ["0.50", "1.75x", "approx 2.83x", "1.75x"],
            ["0.30", "2.05x", "approx 6.09x", "approx 2.32x"],
            ["0.10", "2.35x", "31.6x raw (capped to 3.0)", "approx 2.47x"],
            ["0.00", "2.50x", "(undefined - clipped at 0.01)", "approx 2.49x"],
        ],
    )
    _add_callout(
        doc,
        "Why three methods rather than 'pick one'",
        "Different products and segments exhibit different selection-bias shapes. Premium products "
        "with high acceptance saturate (sigmoid); subprime products with very long rejection tails "
        "do not (power); products with stable acceptance behave linearly. Forcing a single "
        "functional form on every segment would misestimate risk in at least some of them. The "
        "platform makes the choice explicit and reviewable rather than hardcoded.",
    )
    _add_heading(doc, "8.3 Bayesian smoothing", 2)
    _add_paragraph(
        doc,
        "Per-bin acceptance rates are noisy when sample size is small. A Beta-Binomial posterior "
        "shrinks each bin's rate toward the global rate with a configurable prior strength:",
    )
    _add_formula(
        doc,
        "smoothed_rate = (n_booked + alpha) / (n_total + alpha + beta)",
    )
    _add_paragraph(
        doc,
        "where alpha = prior_strength * global_rate and beta = prior_strength * (1 - global_rate). "
        "The reject_bayesian_prior_strength parameter (default 10.0) controls shrinkage:",
    )
    _add_table(
        doc,
        ["Prior strength", "Effective behaviour", "When to use"],
        [
            ["1 - 5", "Minimal shrinkage; per-bin rates barely move", "Rich data, large per-bin counts (>= 100)"],
            [
                "10 - 50 (default range)",
                "Stabilises bins with fewer than 50 observations toward the global rate",
                "Standard / default for most segments",
            ],
            [
                "100+",
                "Pulls all bins strongly toward the global rate",
                "Sparse segments where per-bin rates are not credible",
            ],
        ],
    )
    _add_paragraph(
        doc,
        "Bayesian shrinkage is preferred over ad-hoc count thresholds because the strength of "
        "shrinkage is a single auditable parameter rather than a cliff at, say, n = 30. It also "
        "degrades gracefully: a bin with n = 10 is shrunk substantially; a bin with n = 1,000 is "
        "barely touched. The same epsilon-handling convention is used in PSI/CSI for consistency.",
    )
    _add_heading(doc, "8.4 Marginal monotonicity enforcement on the multipliers", 2)
    _add_paragraph(
        doc,
        "Multipliers are post-processed with isotonic regression (sklearn.isotonic.IsotonicRegression) "
        "via alternating projections across axes so that the implied risk uplift is non-decreasing "
        "along each variable's axis in the expected risk direction. The procedure: for each "
        "variable, average multipliers across the other axes, fit isotonic regression along that "
        "variable's sorted bins, map back to cells, repeat until convergence (max change < 1e-6 or "
        "10 iterations). This guarantees the corrected surface respects the business prior even "
        "when raw bin estimates are noisy.",
    )
    _add_heading(doc, "8.5 Hard guardrails", 2)
    _add_table(
        doc,
        ["Guardrail", "Behaviour", "Reason"],
        [
            [
                "Floor",
                "Multiplier >= 1.0",
                "Risk is never adjusted downward by reject inference; the correction is one-sided by construction (the rejected population should not be safer than the booked one in the same bin).",
            ],
            [
                "Cap",
                "Multiplier <= reject_max_risk_multiplier (default 3.0; configurable up to 10.0)",
                "Prevents runaway uplift in extremely low-acceptance bins where the rate is unreliable. A cap of 3.0 means risk can at most triple — a strong but not absurd correction.",
            ],
            [
                "No / low-demand bins",
                "Shrink toward a conservative low anchor — the reject_no_demand_anchor_percentile (default 10th) percentile of observed rates, floored at 0.01 (0.05 if none observable) — by the confidence weight 1 - exp(-n / reject_confidence_scale)",
                "Avoids NaN propagation and stays conservative: no-demand bins collapse to the low anchor (high uplift), well-observed bins are ~unchanged. Replaces the old anti-conservative median fill.",
            ],
            [
                "Confidence diagnostics",
                "Per-bin effective sample size (ri_bin_count_effective) and confidence score (1 - exp(-n / reject_confidence_scale), default scale 10)",
                "Surfaces low-confidence bins to reviewers without using them as MILP coefficients (separation of concerns).",
            ],
        ],
    )
    _add_paragraph(doc, "Confidence-score reference values (at the default reject_confidence_scale = 10):")
    _add_table(
        doc,
        ["Effective observations", "Confidence score"],
        [
            ["5", "0.39"],
            ["10", "0.63"],
            ["20", "0.86"],
            ["30", "0.95"],
            ["50", "0.99"],
        ],
    )
    _add_heading(doc, "8.6 Optional temporal weighting", 2)
    _add_paragraph(
        doc,
        "Acceptance rates can drift over time as the upstream score model is recalibrated or as "
        "policy changes. Two optional controls localise the rate to recent behaviour:",
    )
    _add_bullets(
        doc,
        [
            "reject_acceptance_recent_months: hard window — only the last N months count.",
            "reject_acceptance_decay_half_life_months: exponential decay — older observations contribute less, with the configured half-life.",
        ],
    )
    _add_paragraph(
        doc,
        "Both default to off (use the full main-period window). The half-life decay is the "
        "preferred form when a window is needed because it avoids cliff effects at the window "
        "boundary.",
    )
    _add_heading(doc, "8.7 Automated parameter tuning (RI optimizer)", 2)
    _add_paragraph(
        doc,
        "The reject_uplift_factor and reject_max_risk_multiplier parameters can be tuned "
        "automatically by minimising the exposure-weighted squared relative error against a "
        "calibration target derived from the standard 1/a^gamma selection-bias model:",
    )
    _add_formula(doc, "target_risk_bin = booked_risk_bin / max(acceptance_rate_bin, 0.05) ^ gamma")
    _add_formula(
        doc,
        "Error = sum(w_i x ((predicted_risk_i - target_risk_i) / target_risk_i)^2) / sum(w_i)",
    )
    _add_paragraph(
        doc,
        "where w_i = todu_amt_pile_h6 (exposure weight). The gamma exponent is configurable:",
    )
    _add_table(
        doc,
        ["gamma value", "Behaviour", "When to choose"],
        [
            ["1.0", "Standard 1/a model: full correction for selection bias", "Default; theoretically grounded"],
            [
                "0.7",
                "Moderate correction, assumes partial sorting within bins",
                "When bins are wide and within-bin heterogeneity is substantial",
            ],
            [
                "0.5",
                "Conservative correction, appropriate when bins are coarse",
                "Few bins per dimension, or score model is known to be weak",
            ],
            ["-> 0", "Minimal correction, target approaches booked risk", "Diagnostic / sanity-check only"],
        ],
    )
    _add_paragraph(doc, "Two search strategies:")
    _add_bullets(
        doc,
        [
            "Grid search (default): exhaustive over an 11 x 9 = 99 combination grid. Deterministic and transparent — preferred when reviewers want to see every evaluated point in ri_optimizer_results.csv.",
            "Optuna TPE: continuous parameter ranges with seeded TPE sampling and 100 trials by default. More sample-efficient when the search space is widened.",
        ],
    )
    _add_paragraph(
        doc,
        "Selection rule: among feasible parameter pairs (a pair is feasible iff it produces a "
        "valid MILP solution at optimum_risk), select the one with lowest calibration error; ties "
        "within 5% are broken by maximising production. This favours equally-well-calibrated "
        "parameters that are less restrictive.",
    )
    _add_paragraph(
        doc,
        "Best parameters are then validated on the MR period and a degradation ratio is reported:",
    )
    _add_formula(doc, "degradation_ratio = mr_calibration_error / main_calibration_error")
    _add_table(
        doc,
        ["Degradation ratio", "Interpretation", "Action"],
        [
            ["~ 1.0", "Stable: RI correction generalises well", "Proceed with confidence"],
            ["1.0 - 2.0", "Moderate degradation: some temporal drift", "Investigate; consider lower gamma"],
            [
                "> 2.0",
                "Significant: RI parameters may be overfit",
                "Lower gamma, coarsen bins, or revert to manual values",
            ],
        ],
    )

    # ====================================================================
    # 9. MILP OPTIMIZATION
    # ====================================================================
    _add_heading(doc, "9. The MILP Optimization", 1)
    _add_paragraph(
        doc,
        "Once every grid cell has an estimate of production, risk numerator, and risk denominator, "
        "the platform solves a Mixed-Integer Linear Program to choose the acceptance mask. The "
        "CellGrid class normalises the aggregated data into a problem where each unique combination "
        "of bins is a 'cell' with a binary decision variable.",
    )
    _add_heading(doc, "9.1 Formal problem statement", 2)
    _add_paragraph(doc, "Decision variables:")
    _add_formula(doc, "x[i] in {0, 1}  for each cell i in the N-dimensional grid")
    _add_paragraph(doc, "where x[i] = 1 means all applications in cell i are accepted.")
    _add_paragraph(doc, "Objective (maximise production):")
    _add_formula(doc, "maximize  sum_i  oa_amt_h0[i] * x[i]")
    _add_paragraph(
        doc,
        "Risk constraint (linearised). The true constraint is the ratio "
        "multiplier x sum(num x x) / sum(den x x) <= target / 100, which is non-linear. Because the "
        "multiplier is a constant scalar, the inequality can be rewritten in linear form:",
    )
    _add_formula(
        doc,
        "sum_i (multiplier x num[i]  -  (target/100) x den[i]) x x[i]  <=  0",
    )
    _add_paragraph(
        doc,
        "Each cell contributes a coefficient that is positive when the cell's risk exceeds the "
        "target and negative when below. The linearisation is exact because the multiplier is "
        "constant by construction (it is the annualisation factor).",
    )
    _add_heading(doc, "9.2 Why MILP rather than a heuristic or a single threshold", 2)
    _add_bullets(
        doc,
        [
            "Globally optimal for the formulated problem: no risk of being trapped in a local maximum that a greedy heuristic might find. A greedy 'add the highest-production cell that still fits the budget' rule is provably suboptimal in the presence of monotonicity constraints.",
            "Constraints are first-class citizens: risk budget, monotonicity, swap-in caps, and segment floors are all linear inequalities the solver enforces exactly. Adding a new constraint means adding a row to the constraint matrix, not redesigning the algorithm.",
            "Reproducible and auditable: the problem matrix is the same every run; the same inputs produce the same mask. scipy.optimize.milp uses HiGHS, which is deterministic.",
            "Linearisation of the risk ratio is exact: because the multiplier is a constant, the ratio constraint becomes a linear inequality without approximation.",
            "Ratio objectives (e.g. 'maximise production for any risk') are handled by the Pareto sweep across multiple risk targets, not by quadratic or fractional programming — keeping the problem solvable with off-the-shelf MILP libraries.",
        ],
    )
    _add_heading(doc, "9.3 Monotonicity constraints", 2)
    _add_paragraph(
        doc,
        "For each pair of adjacent cells along each dimension, the platform adds a constraint:",
    )
    _add_formula(doc, "x[riskier_cell] - x[safer_cell] <= 0")
    _add_paragraph(
        doc,
        "This ensures: if a safer cell is rejected, all riskier cells along that dimension must "
        "also be rejected. The result is a 'staircase' acceptance pattern. Monotonicity is "
        "enforced marginally (per dimension independently), not jointly across dimensions.",
    )
    _add_paragraph(
        doc,
        "Why marginal monotonicity rather than joint monotonicity:",
    )
    _add_bullets(
        doc,
        [
            "Computationally cheap: O(N_cells x N_dims) constraints rather than O(N_cells^2). For a 10x10x3 grid, that is 540 constraints vs. up to 90,000.",
            "Sufficient in practice: grid connectivity tends to prevent jointly non-monotone patterns from being optimal, even though they are technically feasible.",
            "Operationally meaningful: a policy that is monotone per dimension is interpretable as a per-axis cutoff; a jointly monotone policy that is not marginally monotone has no simple verbal description.",
        ],
    )
    _add_heading(doc, "9.4 Why monotonicity constraints at all", 2)
    _add_paragraph(
        doc,
        "Without monotonicity, the MILP could legally accept a risky cell while rejecting a safer "
        "neighbouring cell — a 'swiss cheese' policy that no underwriting committee would approve. "
        "Monotonicity:",
    )
    _add_bullets(
        doc,
        [
            "Is explainable to the front line: 'we accept everything to the left of and below this line per axis'.",
            "Resists adverse selection: applicants cannot game the policy by aiming for a safer-looking cell that is rejected.",
            "Aligns with regulatory expectations of consistent risk treatment across comparable applicants — a borderline applicant in a safer cell should never be rejected when a rougher-profile applicant in a riskier cell is accepted.",
            "Is a smoothing prior: it reduces overfitting to per-cell noise by forcing the policy to follow the underlying risk gradient.",
        ],
    )
    _add_heading(doc, "9.5 Swap-in constraints", 2)
    _add_paragraph(
        doc,
        "Two optional constraints cap the contribution of repesca to the optimised portfolio:",
    )
    _add_paragraph(doc, "Production-share cap (max_swapin_production_pct):")
    _add_formula(
        doc,
        "sum_i (oa_amt_h0_rep[i]  -  (pct/100) x oa_amt_h0[i]) x x[i]  <=  0",
    )
    _add_paragraph(doc, "Swap-in risk cap (max_swapin_risk):")
    _add_formula(
        doc,
        "sum_i (multiplier x num_rep[i]  -  (max_risk/100) x den_rep[i]) x x[i]  <=  0",
    )
    _add_paragraph(
        doc,
        "These let management bound the policy's exposure to the most uncertain population — where "
        "reject-inference assumptions carry the most weight — without abandoning the optimisation. "
        "When None, the MILP behaves as before. If too tight, the solver returns infeasible and the "
        "platform falls back (Section 9.7).",
    )
    _add_heading(doc, "9.6 Fixed-cell constraints", 2)
    _add_paragraph(
        doc,
        "Individual cells can be pinned as forced-accept or forced-reject before re-optimisation: "
        "the solver sets lb = ub = 1 (forced-accept) or lb = ub = 0 (forced-reject) for the pinned "
        "cells while monotonicity and risk constraints remain active. If the pins contradict each "
        "other or the risk budget, the solver returns None — surfacing the conflict rather than "
        "silently dropping pins.",
    )
    _add_paragraph(
        doc,
        "This supports two governance use cases: (a) a risk committee mandates that certain cells "
        "must always be accepted (or rejected) for policy reasons, and the optimiser respects "
        "those constraints; (b) the dashboard's Pin Mode lets analysts cycle cells through "
        "unpinned -> accept -> reject and re-run the MILP interactively to see the impact.",
    )
    _add_heading(doc, "9.7 Solver and fallbacks", 2)
    _add_paragraph(
        doc,
        "The MILP is solved by scipy.optimize.milp (HiGHS backend) with a configurable time limit "
        "(milp_time_limit, default 30 seconds). If the solver returns infeasible — typically because "
        "of an aggressive risk target combined with tight swap-in caps — the platform falls back:",
    )
    _add_table(
        doc,
        ["Case", "Fallback", "Why"],
        [
            [
                "N > 2 dimensions, infeasible",
                "Genetic algorithm via pymoo (NSGA-II)",
                "Handles large combinatorial spaces; produces a near-optimal Pareto frontier rather than a single point.",
            ],
            [
                "N = 2 dimensions, infeasible",
                "Legacy enumeration over all monotonic combinations",
                "The 2D space is small enough to brute-force; produces the exact frontier.",
            ],
            ["MILP returns optimal", "Use directly", "Standard path"],
        ],
    )
    _add_paragraph(
        doc,
        "Every fallback is logged so reviewers can see when and why the exact solver did not "
        "return a solution. Fallback frequency is itself a methodology indicator: a segment that "
        "frequently falls back is signalling that its constraint set is over-binding.",
    )

    # ====================================================================
    # 10. PARETO FRONTIER AND SCENARIOS
    # ====================================================================
    _add_heading(doc, "10. Pareto Frontier and Scenario Analysis", 1)
    _add_paragraph(
        doc,
        "Rather than committing to a single risk target, the platform sweeps the MILP across "
        "pareto_n_points (default 50) evenly-spaced risk targets from 0.01% to "
        "max_risk x 1.1, where max_risk is the b2_ever_h6 obtained by accepting all cells. "
        "Solutions are deduplicated by acceptance mask and filtered to the Pareto frontier:",
    )
    _add_numbered(
        doc,
        [
            "Monotone production filter: keep solution i only if its production strictly exceeds every solution with lower risk.",
        ],
    )
    _add_paragraph(
        doc,
        "For the standard two-objective case (risk vs production) this single sort-and-sweep is "
        "already exact, so the separate full-dominance pass — removing any solution with "
        "lower-or-equal risk and higher-or-equal production (with at least one strict inequality) — "
        "is intentionally skipped; it would not change the frontier. The resulting frontier has "
        "strictly increasing production as risk increases — a clean, monotone trade-off curve.",
    )
    _add_heading(doc, "10.1 Why a frontier rather than a point", 2)
    _add_bullets(
        doc,
        [
            "Separation of concerns: the technical decision (efficient policies) is delivered by analytics; the business decision (risk appetite) belongs to the committee. The frontier delivers the menu, not the choice.",
            "Concrete cost of risk reduction: management can read 'tightening the risk target by 10 bps costs X EUR of production', which is far more informative than a single recommended cutoff.",
            "Composable across segments: each segment's frontier feeds the global allocation MILP, which distributes a portfolio-wide risk budget by selecting one frontier point per segment to maximise total production. This decouples segment-level optimisation from portfolio-level governance.",
            "Robustness signal: a frontier with many densely-packed points indicates a smooth trade-off; one with large jumps signals discrete cliff effects worth investigating before deployment.",
        ],
    )
    _add_heading(doc, "10.2 Scenario selection", 2)
    _add_paragraph(
        doc,
        "From the frontier the platform highlights three operating points around the configured "
        "optimum_risk (e.g. 1.1%) with risk_step (e.g. 0.1):",
    )
    _add_table(
        doc,
        ["Scenario", "Risk target", "Purpose"],
        [
            [
                "Pessimistic",
                "optimum_risk - risk_step",
                "Acceptance policy under a tightened risk appetite — quantifies upside foregone if risk view worsens",
            ],
            ["Base", "optimum_risk", "Recommended operating point"],
            [
                "Optimistic",
                "optimum_risk + risk_step",
                "Policy under a loosened appetite — quantifies upside available if risk view improves",
            ],
        ],
    )
    _add_paragraph(
        doc,
        "For each scenario the same downstream pack is generated: bootstrap CIs (1,000 resamples), "
        "MR validation, PSI/CSI, audit table, sensitivity outputs. This forces every recommendation "
        "to be considered across at least three risk environments before deployment.",
    )
    _add_heading(doc, "10.3 Bootstrap confidence intervals", 2)
    _add_paragraph(
        doc,
        "For each selected solution, the pipeline runs 1,000 bootstrap resamples of the booked "
        "population, recalculating production and risk for each sample. The 2.5th and 97.5th "
        "percentiles (np.nanpercentile, NaN-aware) provide 95% confidence intervals. NaN-valued "
        "replicates (from empty-denominator bins) are excluded from the percentile computation "
        "rather than corrupting the CI. The resulting CIs are reported alongside point estimates "
        "in every output table.",
    )

    # ====================================================================
    # 11. MR VALIDATION
    # ====================================================================
    _add_heading(doc, "11. Out-of-Time Validation (MR Period)", 1)
    _add_paragraph(
        doc,
        "Models that look excellent in-sample routinely fail out-of-sample. The MR period is a "
        "holdout block of recent data not used in optimisation; cutoffs are validated on it before "
        "they are recommended to committee.",
    )
    _add_heading(doc, "11.1 Risk-source priority per bin", 2)
    _add_paragraph(
        doc,
        "MR risk is sourced bin-by-bin from the most reliable evidence available, in this priority "
        "order. The chosen source is logged in the risk_source column for every bin so reviewers "
        "can see exactly which evidence drove every estimate:",
    )
    _add_table(
        doc,
        ["Priority", "Source", "Condition", "Rationale"],
        [
            [
                "1",
                "mr_observed",
                "n_obs_mr >= mr_min_obs_per_bin (default 30) mature MR-period H6",
                "Direct MR-period H6 — the highest-quality evidence; observed mature H6 takes precedence over H3 extrapolation",
            ],
            [
                "2",
                "h3_extrapolated",
                "MR H6 insufficient, but mature MR H3 exists (n_obs_mr_h3 >= max(mr_min_obs_per_bin // 2, 10), the relaxed H3 gate) and the main-period H6/H3 ratio is finite",
                "Scales MR-observed H3 by the main-period H6/H3 ratio; fills bins where direct H6 is still too thin",
            ],
            [
                "3",
                "main_imputed",
                "Main-period bin exists but MR evidence is insufficient",
                "Preserves baseline knowledge; the model has been validated on this data",
            ],
            [
                "4",
                "model_fallback",
                "Bin absent from the main period (b2_main is NaN) with no usable MR H6 or H3",
                "Last resort: inferred via the trained risk model",
            ],
        ],
    )
    _add_heading(doc, "11.2 H3 to H6 extrapolation", 2)
    _add_paragraph(
        doc,
        "The 6-month metric needs a 6-month observation window, so the most recent applications "
        "are not yet mature. The 3-month metric matures sooner and is highly correlated with H6. "
        "The platform extrapolates H3 to H6 using the main-period H6/H3 ratio with a configurable "
        "functional form:",
    )
    _add_table(
        doc,
        ["Method", "Formula", "When to use"],
        [
            [
                "linear (default)",
                "b2_mr_h3 x (b2_main / b2_main_h3)",
                "Proportional H6/H3 relationship — the safest default",
            ],
            [
                "power",
                "b2_mr_h3 x ratio x (b2_mr_h3 / b2_main_h3)^(alpha-1)",
                "Convex (alpha > 1) or concave (alpha < 1) departures from linearity; falls back to linear per bin when b2_main_h3 is unavailable",
            ],
            [
                "logistic",
                "b2_mr_h3 x (1 + 2 x tanh(k x (ratio-1)/2) / k)",
                "Caps extreme ratios smoothly when the H6/H3 relationship saturates",
            ],
            [
                "auto (recommended)",
                "Fits curvature alpha from main-period data via weighted log-log regression",
                "Recommended — removes a manual tuning knob and replaces it with an evidence-based choice",
            ],
        ],
    )
    _add_paragraph(doc, "Auto-calibration procedure (method = 'auto'):")
    _add_numbered(
        doc,
        [
            "Run a weighted log-log regression on main-period bins: log(b2_h6) = c + alpha x log(b2_h3).",
            "If the 95% CI for alpha includes 1.0, fall back to linear.",
            "Otherwise use power with alpha clipped to [0.3, 3.0] to prevent absurd extrapolation.",
            "If fewer than 4 valid bins exist, fall back to linear regardless.",
        ],
    )
    _add_paragraph(doc, "Worked example (octroi=2, efx=5):")
    _add_table(
        doc,
        ["Period", "Inputs", "Computed value"],
        [
            ["Main (mature)", "todu_30ever_h6=100, todu_amt_pile_h6=1000, mult=7", "b2_main = 7 x 100/1000 = 70%"],
            ["Main (mature)", "todu_30ever_h3=45, todu_amt_pile_h3=900, mult_h3=4", "b2_main_h3 = 4 x 45/900 = 20%"],
            ["Main", "Ratio", "70 / 20 = 3.5"],
            ["MR (H3 mature, H6 not)", "todu_30ever_h3=10, todu_amt_pile_h3=160", "b2_mr_h3 = 4 x 10/160 = 25%"],
            ["MR linear extrapolation", "25% x 3.5", "87.5%"],
            ["MR power, alpha=1.3", "25% x 3.5 x (25/20)^0.3", "approx 93.6%"],
        ],
    )
    _add_paragraph(
        doc,
        "Safeguards: bins where b2_main_h3 is approximately zero skip extrapolation; for power, "
        "bins where b2_main_h3 is NaN automatically fall back to linear; mr_extrapolation_risk_"
        "multiplier (default 3.0) caps extrapolated risk relative to main; mr_extrapolation_hard_"
        "cap (default 15%) caps absolute extrapolated risk.",
    )
    _add_heading(doc, "11.3 Risk production summary on MR", 2)
    _add_paragraph(
        doc,
        "The function calculate_metrics_from_cuts applies the optimal cutoff to MR data and "
        "produces a four-row summary that is the central management deliverable:",
    )
    _add_table(
        doc,
        ["Row", "Origin", "Interpretation"],
        [
            ["Actual (booked)", "Observed booked under current policy", "Current-policy baseline"],
            ["Swap-in (repesca accepted)", "Repesca passing the proposed cutoff", "Upside from opening cutoffs"],
            ["Swap-out (booked rejected)", "Booked failing the proposed cutoff", "Downside from tightening cutoffs"],
            ["Optimum", "Net of the proposed cutoff", "Expected outcome under the new policy"],
        ],
    )
    _add_paragraph(
        doc,
        "When H3 data is available, each row also includes Risk H3 (%) — the observed 3-month risk "
        "metric — providing an early-risk view alongside the standard H6 metric. Both H6 and H3 "
        "appear in the HTML consolidated report and the Excel workbook (RP MR sheets).",
    )
    _add_callout(
        doc,
        "Why this matters for review",
        "MR validation is the platform's principal external evidence that the optimisation will "
        "work in production. A degradation between main-period and MR-period calibration error is "
        "the single strongest signal that the methodology has overfit the historical window — and "
        "the platform reports this ratio explicitly for every RI optimisation run.",
    )

    # ====================================================================
    # 12. STABILITY
    # ====================================================================
    _add_heading(doc, "12. Stability Analysis (PSI / CSI)", 1)
    _add_paragraph(
        doc,
        "PSI (Population Stability Index) and CSI (Characteristic Stability Index) quantify how "
        "much the distribution of inputs has shifted between the optimisation period and the MR "
        "period. The platform uses the standard formula:",
    )
    _add_formula(
        doc,
        "PSI = sum (Actual%_i - Expected%_i) x ln(Actual%_i / Expected%_i)",
    )
    _add_paragraph(
        doc,
        "A unified epsilon constant (PSI_EPSILON = 0.0001) is applied only inside the log() term "
        "for zero-percentage bins, while the difference term (Actual% - Expected%) uses the "
        "original unmodified proportions. This preserves the standard PSI scale and ensures "
        "non-negativity. Both PSI and CSI use the same epsilon application method (.where()) for "
        "consistency.",
    )
    _add_table(
        doc,
        ["PSI range", "Interpretation", "Action"],
        [
            ["< 0.10", "Stable — no significant change", "Proceed with confidence"],
            ["0.10 - 0.25", "Moderate drift — investigation recommended", "Investigate cause before changing policy"],
            [">= 0.25", "Unstable — action required", "Hold or roll back; refresh model and rerun"],
        ],
    )
    _add_paragraph(
        doc,
        "PSI is computed per binned variable. CSI uses the same formula applied to categorical "
        "distributions (e.g. segment composition). Both are written to stability_psi_{scenario}.csv "
        "and visualised in stability_report_{scenario}.html.",
    )
    _add_heading(doc, "12.1 Why this construction of PSI and CSI", 2)
    _add_bullets(
        doc,
        [
            "Industry-standard formulation: any reviewer with credit-risk experience can interpret the values directly.",
            "Single epsilon location (inside log only): preserves the conventional scale and avoids the slight bias that comes from epsilon-padding both terms.",
            "Where()-based application: handles zero bins safely without distorting non-zero bins.",
            "Per-variable computation: surfaces which variable is driving instability rather than reporting a single aggregate number that hides the cause.",
        ],
    )

    # ====================================================================
    # 13. SENSITIVITY
    # ====================================================================
    _add_heading(doc, "13. Sensitivity and Marginal-Impact Analysis", 1)
    _add_heading(doc, "13.1 Risk perturbation", 2)
    _add_paragraph(
        doc,
        "Sensitivity analysis perturbs the predicted risk column (todu_30ever_h6_rep) by "
        "+/- 5%, 10%, 20% (configurable via sensitivity_levels) and re-solves the MILP at each "
        "level, respecting active swap-in constraints and the MILP time limit. For each perturbation "
        "level, the platform records the new optimum, the cells that flipped, and the production / "
        "risk delta versus baseline.",
    )
    _add_heading(doc, "13.2 Cell flip thresholds", 2)
    _add_paragraph(
        doc,
        "For each cell, the platform identifies the smallest perturbation that flips its accept / "
        "reject status. This produces a per-cell 'fragility' score: cells that flip at +/- 5% are "
        "borderline; cells that survive +/- 20% are robust. Output is written to "
        "sensitivity_cell_detail.csv.",
    )
    _add_heading(doc, "13.3 Marginal impact (analytical, O(N))", 2)
    _add_paragraph(
        doc,
        "Marginal-impact analysis analytically computes the production and risk delta of flipping "
        "each cell — O(N) total cost. Baseline portfolio sums are computed once, then per-cell "
        "deltas are evaluated by simple arithmetic. Output (cell_marginal_impact.csv) reports "
        "delta_production, delta_risk_pct, cell_production, and cell_risk for every cell.",
    )
    _add_paragraph(
        doc,
        "Together, these three outputs answer the management question 'how robust is this policy?' "
        "in concrete terms: which cells are on the edge, what would tip them, and what would the "
        "portfolio look like if they tipped.",
    )

    # ====================================================================
    # 14. TRENDS
    # ====================================================================
    _add_heading(doc, "14. Trend Monitoring and Anomaly Detection", 1)
    _add_paragraph(
        doc,
        "Monthly metrics (approval rate, production, risk) are tracked over time with Statistical "
        "Process Control. The control band uses:",
    )
    _add_bullets(
        doc,
        [
            "A one-period-lagged rolling median as the centre line — lagging prevents the current observation from biasing the band that is supposed to detect it.",
            "A robust moving-range (I-MR) scale — sigma is estimated from the median of the absolute month-to-month differences (median(MR) / 0.9539), i.e. from within-subgroup short-term variation rather than the dispersion of the levels. This keeps the band tight under a slow drift, so a point that jumps off the trend is still flagged instead of being masked by an inflated spread.",
            "A plain Shewhart n-sigma multiplier (3-sigma by default) — stable at any window size, with no small-sample critical-value inflation that would over-widen the band.",
        ],
    )
    _add_paragraph(
        doc,
        "Anomalies are flagged when a metric crosses the band. Output is written to "
        "trend_anomalies_{segment}.csv and visualised in metric_trends_{segment}.html.",
    )

    # ====================================================================
    # 15. PORTFOLIO ALLOCATION
    # ====================================================================
    _add_heading(doc, "15. Multi-Segment Portfolio Allocation", 1)
    _add_paragraph(
        doc,
        "After per-segment optimisation, the platform can allocate a global risk budget across "
        "segments by selecting one point from each segment's efficient frontier to maximise total "
        "production. This solves the governance problem: 'we have a portfolio-wide risk appetite, "
        "but the segments don't all need the same risk level — how should we distribute?'",
    )
    _add_table(
        doc,
        ["Method", "How it works", "When to use"],
        [
            [
                "exact (MILP)",
                "Global MILP via scipy.optimize.milp; selects one frontier point per segment subject to global risk and production constraints",
                "Default; produces the provably optimal allocation",
            ],
            [
                "greedy",
                "Hill-climbing heuristic; faster but may find local optima",
                "Diagnostic; or when the exact solver returns infeasible",
            ],
        ],
    )
    _add_paragraph(
        doc,
        "Per-segment constraints (min_risk, max_risk, min_production, locked_sol_fac) are enforced "
        "as MILP inequalities. The exact solver falls back to greedy if infeasible. Outputs include "
        "a per-segment allocation table, a policy/cutoff table for portfolio owners, and a "
        "Markdown narrative explaining which constraints bound the solution — written for the risk "
        "committee, not for analysts.",
    )

    # ====================================================================
    # 16. CONTROLS
    # ====================================================================
    _add_heading(doc, "16. Controls, Auditability, and Reproducibility", 1)
    _add_table(
        doc,
        ["Control", "Mechanism", "Evidence"],
        [
            [
                "Versioned configuration",
                "TOML files validated by Pydantic (PreprocessingSettings)",
                "config.toml + segments.toml committed alongside outputs",
            ],
            [
                "Idempotent runs",
                "Seeded randomness (numpy 42, Optuna 42); deterministic MILP; deterministic CV folds",
                "Re-running with the same TOML produces byte-identical outputs",
            ],
            [
                "Decision logging",
                "risk_source per MR bin; ri_optimizer_results.csv per RI run; fallback events logged at WARNING",
                "Every methodological decision is traceable to its inputs",
            ],
            [
                "Quality gates",
                "DQ checks (schema, missing values, outliers, date ranges, categorical consistency) before any modelling",
                "Failed checks block the run unless --skip-dq-checks is explicitly passed",
            ],
            [
                "Test coverage",
                "pytest with synthetic DataFrames; coverage target 80% on changed code",
                "CI gates lint -> test -> Docker build on push/PR",
            ],
            [
                "Immutable outputs",
                "Dated, structured output/{segment}/ folders; HTML report bundles all assets self-contained",
                "Outputs can be archived without losing context",
            ],
            [
                "Evidence pack ships with every recommendation",
                "MR validation, PSI/CSI, bootstrap CIs, sensitivity, marginal impact are first-class outputs (not optional)",
                "No recommendation can be made without its supporting evidence",
            ],
        ],
    )
    _add_callout(
        doc,
        "What an auditor can do with a finished run",
        "Open the run's TOML, read the bin definitions and parameter choices; open ri_optimizer_"
        "results.csv to see every parameter pair evaluated and the chosen one; open mr_risk_"
        "comparison_{scenario}.csv and inspect risk_source for each bin; replicate the run by "
        "re-executing main.py against the same TOML — outputs will match byte-for-byte. There is no "
        "step in the pipeline that depends on uncaptured operator judgement.",
    )

    # ====================================================================
    # 17. ASSUMPTIONS AND LIMITATIONS
    # ====================================================================
    _add_heading(doc, "17. Assumptions and Limitations", 1)
    _add_paragraph(
        doc,
        "The platform is explicit about its assumptions so reviewers can check them. The most material are:",
    )
    _add_table(
        doc,
        ["#", "Assumption", "Why it might fail", "Mitigation in the platform"],
        [
            [
                "1",
                "Score discriminance is reasonably stable across the optimisation and MR windows",
                "Material score-engine changes invalidate the trained risk surface",
                "PSI / CSI checks; MR degradation ratio; explicit notification when score recalibration is detected",
            ],
            [
                "2",
                "Reject inference: within a bin, accepted and rejected applicants differ in expected risk in a way captured by the parceling model",
                "Within-bin heterogeneity means the correction is approximate",
                "Three functional forms (linear/power/sigmoid); Bayesian smoothing for sparse bins; out-of-time validation of RI parameters with degradation ratio",
            ],
            [
                "3",
                "Macroeconomic conditions during the MR period are a reasonable proxy for the policy-implementation period",
                "A shift in unemployment, inflation, or interest rates between MR and deployment can invalidate calibrated risk",
                "Through-the-cycle stress is out of scope and must be handled at portfolio-policy level. The platform reports MR drift to surface this.",
            ],
            [
                "4",
                "Marginal monotonicity is enforced (per dimension), not joint",
                "Theoretically allows non-jointly-monotone acceptance patterns",
                "Grid connectivity makes this rare in practice; sensitivity analysis surfaces any such patterns",
            ],
            [
                "5",
                "The risk-budget constraint is linearised by treating the multiplier as a constant",
                "Would fail if the multiplier varied by cell",
                "Exact by construction — multiplier is always a constant (annualisation factor)",
            ],
            [
                "6",
                "Stress factor extrapolates from the worst 5% of booked to the rejected population",
                "Tail of booked is not necessarily representative of rejected risk",
                "Per-bin mode for heterogeneous portfolios; disabled mode when parceling is active to avoid double-counting",
            ],
            [
                "7",
                "tasa_fin (transformation rate) is uniform across bins by default",
                "Conversion rates can vary across cells",
                "per_bin_tasa_fin = true switches to per-cell rates with fallback to global for sparse bins (< 10 records)",
            ],
            [
                "8",
                "The risk model trained on booked data extrapolates to rejected data",
                "The risk-score relationship may differ between populations",
                "Stress factor + reject-inference parceling correct for this; monotonic constraints prevent extreme extrapolation",
            ],
            [
                "9",
                "b2_ever_h6 is a sufficient leading indicator of loss",
                "Pricing, recovery, and lifetime PD are not modelled",
                "Out of scope; this platform feeds origination decisions, not provisioning models",
            ],
        ],
    )

    # ====================================================================
    # 18. DEFAULT PARAMETERS
    # ====================================================================
    _add_heading(doc, "18. Default Parameter Reference", 1)
    _add_paragraph(
        doc,
        "For management reference, the most commonly inspected defaults are:",
    )
    _add_table(
        doc,
        ["Parameter", "Default", "Range", "Rationale for the default"],
        [
            [
                "multiplier (H6)",
                "7.0",
                "> 0",
                "Annualisation constant — calibrated to convert 6-month observations to 12-month equivalents",
            ],
            ["multiplier (H3)", "4.0", "> 0", "H3 annualisation constant"],
            [
                "optimum_risk",
                "1.1%",
                "Per-segment",
                "Default risk appetite; should be set per segment in segments.toml",
            ],
            ["risk_step", "0.1", "—", "Pessimistic / optimistic offsets around optimum"],
            [
                "pareto_n_points",
                "50",
                "5 - 500",
                "Dense enough for a smooth frontier; light enough to solve in seconds",
            ],
            ["n_bootstraps", "1000", "100 - 50000", "Standard for stable 95% percentile CIs"],
            [
                "cv_folds",
                "4-5",
                "2 - 10",
                "Standard for model selection; higher folds increase variance, lower folds increase bias",
            ],
            ["milp_time_limit", "30 s", "> 0", "Generous for 2D / 3D problems; falls back if exceeded"],
            [
                "sensitivity_levels",
                "+/- 5, 10, 20%",
                "—",
                "Symmetric perturbation grid covering near-policy and tail scenarios",
            ],
            ["mr_min_obs_per_bin", "30", ">= 1", "Statistical rule of thumb for stable per-bin estimates"],
            [
                "stress_mode",
                "global",
                "global / per_bin / disabled",
                "Conservative default; switch to per_bin when score variables are heterogeneous",
            ],
            [
                "reject_uplift_factor",
                "1.5",
                "0 - 10",
                "Moderate default for parceling; tune per segment via RI optimizer",
            ],
            ["reject_max_risk_multiplier", "3.0", "1 - 10", "Risk can at most triple — strong but not absurd"],
            ["reject_bayesian_prior_strength", "10.0", "0 - 1000", "Stabilises bins with < 50 observations"],
            [
                "mr_extrapolation_method",
                "linear (auto recommended)",
                "linear / power / logistic / auto",
                "Auto removes a manual tuning knob",
            ],
            ["mr_extrapolation_risk_multiplier", "3.0", "(0, 10]", "Caps extrapolated risk relative to main-period"],
            ["mr_extrapolation_hard_cap", "15%", "(0, 100]", "Absolute cap on extrapolated risk"],
            ["PSI_EPSILON", "0.0001", "—", "Standard epsilon for log-term zero-bin handling"],
        ],
    )

    # ====================================================================
    # 19. GLOSSARY
    # ====================================================================
    _add_heading(doc, "19. Glossary", 1)
    _add_table(
        doc,
        ["Term", "Definition"],
        [
            [
                "b2_ever_h6",
                "Annualised 6-month vintage rate of accounts ever 30+ days past due, weighted by exposure. Primary risk metric.",
            ],
            [
                "b2_ever_h3",
                "3-month equivalent of b2_ever_h6, used for early warning and H3 -> H6 extrapolation when the MR window is not yet mature.",
            ],
            ["Booked", "Application that was approved and disbursed; has observed outcomes."],
            [
                "Repesca / score-rejected",
                "Application rejected by the credit-score cutoff (reject_reason = '09-score'); outcomes unobserved and inferred.",
            ],
            [
                "Policy-rejected",
                "Application rejected for non-score reasons (fraud, KYC, documentation; reject_reason = '08-other'); excluded from optimisation.",
            ],
            ["Cell", "One unique combination of bin coordinates in the N-dimensional score grid."],
            ["Mask", "Binary vector indicating which cells are accepted under a candidate policy."],
            [
                "MILP",
                "Mixed-Integer Linear Program; the solver framework used to choose cutoffs (scipy.optimize.milp / HiGHS).",
            ],
            ["Pareto frontier", "Set of policies for which no other policy is both lower-risk and higher-production."],
            ["MR period", "Recent Monitoring period — out-of-time holdout used for validation, not for optimisation."],
            ["PSI / CSI", "Population / Characteristic Stability Index; measure distribution drift between periods."],
            [
                "Stress factor",
                "Scalar (or per-bin map) that uplifts repesca risk predictions to reflect tail-risk concentration in the booked population.",
            ],
            ["tasa_fin", "Transformation rate; the fraction of eligible applications that are ultimately disbursed."],
            [
                "Swap-in / swap-out",
                "Repesca accepted under a new policy / booked rejected under a new policy — the two delta populations vs. the current cutoff.",
            ],
            [
                "1-SE rule",
                "Selection rule that picks the simplest model within one standard error of the lowest mean CV error — favours stability over chasing minimum mean.",
            ],
            [
                "Annual coefficient",
                "12 / n_months_in_observation_period — normalises flow indicators to a 12-month equivalent.",
            ],
            [
                "Cutoff floor segment",
                "A segment whose accepted cells must be a subset of this segment's accepted cells, enforcing nested cutoffs across segments.",
            ],
            [
                "Supersegment",
                "Group of segments that share a trained model (modelling supersegment) and/or are reported together (reporting supersegment).",
            ],
            [
                "Degradation ratio",
                "MR-period calibration error divided by main-period calibration error; > 2 indicates likely overfit.",
            ],
            [
                "Marginal monotonicity",
                "Monotonicity enforced per dimension independently, not jointly across dimensions.",
            ],
            [
                "Audit table",
                "Record-level classification (keep / swap_in / swap_out / rejected / rejected_other) under a proposed cutoff.",
            ],
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
